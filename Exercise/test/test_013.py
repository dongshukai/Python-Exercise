# -*- coding: utf-8 -*-

import openpyxl
from openpyxl import chart, chartsheet
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Inches, Cm
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.shape import WD_INLINE_SHAPE, WD_INLINE_SHAPE_TYPE
from docx.shape import InlineShape, InlineShapes

# from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

# document = Document()
# # part = document.part
# #
# # paragraph = document.add_paragraph('BACX')
# paragraph = document.add_paragraph()
# run = paragraph.add_run('AAAAA')
# # picture = run.add_picture(u'/Users/dongsk/Downloads/图片/董.jpg', width=Inches(12.5))
# # # paragraph.add_run(WD_INLINE_SHAPE.SMART_ART.denominator)
# M_inline = chart.LineChart
# shape = InlineShape(M_inline)
#
# #
# document.save(u'椭圆3.docx')

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """

    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


docName = '/Users/dongsk/Downloads/payment to hk.docx'
docName = 'payment to hk.docx'
doc = docx.Document(docName)
print doc._body._element.xml
# fulltext = []
# paras = doc.paragraphs
# for p in paras:
#     fulltext.append(p.text)
# M_inline = chart.LineChart
#
# # paras = doc.paragraphs
# # shapes = doc.inline_shapes
# # for p in paras:
# #     print p.text
#

doc1 = docx.Document('空文档.docx')
doc2 = docx.Document('一个椭圆.docx')
doc3 = docx.Document('测试1.docx')
print doc1._body._element.xml
print doc2._body._element.xml
pass




