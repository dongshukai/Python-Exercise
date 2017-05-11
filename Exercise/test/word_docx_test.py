# -*- coding: utf-8 -*-

import xlrd
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.shape import WD_INLINE_SHAPE, WD_INLINE_SHAPE_TYPE
from docx.shape import InlineShape, InlineShapes
from docx.enum.style import WD_BUILTIN_STYLE as WD_STYLE

from datetime import datetime
from cStringIO import StringIO

from docx.oxml import register_element_cls


MONTH_EN_MAP = {
    '01': 'January',
    '02': 'February',
    '03': 'March',
    '04': 'April',
    '05': 'May',
    '06': 'June',
    '07': 'July',
    '08': 'August',
    '09': 'September',
    '10': 'October',
    '11': 'November',
    '12': 'December',
}

KEY_LIST = [u'name', u'value', u'value_type', u'bank_number', u'is_only', u'cumulative_item', u'country', u'Phone']

def set_column_width(columns, width_list=[1.28, 3.25, 2, 4.18, 4.18, 2.57, 4.25, 2.93, 3.07]):
    # 列
    columns_position = [WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER,
                        WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER,
                        WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER]
    for index, column in enumerate(columns):
        width = Cm(width_list[index])
        for cell in column.cells:
            cell.width = width
            paragraphs = cell.paragraphs
            for p in paragraphs:
                p.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER    # WD_ALIGN_PARAGRAPH.CENTER

def set_row_width(rows, width_list=[3.23, 0.56]):
    hight = Cm(1)
    for index, row in enumerate(rows):
        if index == 0:
            row.height = Cm(width_list[0])
            # row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            # for r_cell in row.cells:
            #     r_cell.height = Cm(width_list[0])
        else:
            row.height = Cm(width_list[1])
            # row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            # for r_cell in row.cells:
            #     r_cell.height = Cm(width_list[1])

def set_table_width_attrs(table):
    rows = [3.23, 0.56]
    columns = [1.28, 3.25, 2, 4.18, 4.18, 2.57, 4.25, 2.93, 3.07]
    columns_position = [WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER,
                        WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER,
                        WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.CENTER]


def get_excel_data():
    data = {}
    file_name = '/Users/dongsk/Desktop/其它/工作簿1.xlsx'
    # file_name = '/Users/dongsk/Downloads/1静态数据准备.xlsx'
    file_data = xlrd.open_workbook(file_name)
    tables = file_data.sheets()
    tables = file_data.sheets()
    for table in tables:
        print "开始处理 ",table.name
        nrows = table.nrows    # 行数
        if nrows <= 1:
            print "文件 < %s > 没有内容" % (table.name)
            continue
        ncols = table.ncols     # 列数
        colnames = table.row_values(0)  # 第一行数据
        leave_none_list = []
        data[table.name] = []
        for rownum in range(1, nrows):
            data_dict = {}
            row = table.row_values(rownum)
            if row:
                for i in range(len(colnames)):
                    ccol = colnames[i]
                    rrow = row[i]
                    data_dict[ccol] = rrow
                data[table.name].append(data_dict)
        print table.name, "处理完成"

    return data

def make_docx_file_format(period_year, period_month, total_money, rule_result_data):
    year, month, day = datetime.now().strftime('%Y-%m-%d').split('-')
    document = Document()
    # document.settings.odd_and_even_pages_header_footer = True
    obj_styles = document.styles['Normal']
    # obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_styles.font
    obj_font.size = Pt(10.5)
    obj_font.name = u'宋体'
    # obj_font.name = u'Calibri'
    # 设置页面格式
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(20.99)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.393)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    # document
    # 生成文档头文件
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph_format.keep_with_next = True
    # 1 传真委办发薪交易指示表格Payroll Instruction by Facsimile
    run = paragraph.add_run(u'传真委办发薪交易指示表格')
    # 1.1
    run.font.size = Pt(11)
    run.font.name = u'宋体'
    # 1.2
    run = paragraph.add_run(u'Payroll Instruction by Facsimile')
    run.font.size = Pt(10)
    run.font.name = u'Calibri'
    run.bold = True
    # 2 空行
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    # 3 致：中国银行(香港)有限公司
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    # 3.1
    run = paragraph.add_run(u'致：中国银行(香港)有限公司 ')
    run.font.name=u'Calibri'
    run.font.size = Pt(11)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 4 To: Bank of China (Hong Kong) Limited                                 											日期Date：  30 August 2016
    paragraph = document.add_paragraph()
    # 设置行间距
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    # 4.1
    run = paragraph.add_run(u'To: Bank of China (Hong Kong) Limited')
    run.font.size = Pt(10)
    run.bold = True
    run.font.name = u'Calibri'
    # 4.2
    run = paragraph.add_run(u'                                 											')
    run.font.name = u'Calibri'
    run.bold = True
    run.font.size = Pt(11)
    # 4.3
    run = paragraph.add_run(u'日期')
    run.font.name = u'宋体'
    run.font.size = Pt(11)
    # 4.4
    run = paragraph.add_run(u'Date')
    run.font.name = u'Calibri'
    run.font.size = Pt(11)
    # 4.5
    run = paragraph.add_run(u'：')
    run.font.name = u'宋体'
    run.font.size = Pt(11)
    # 4.6
    run = paragraph.add_run(u'  %s %s %s   ' % (day, MONTH_EN_MAP[str(month)], year))
    run.font.name = u'Calibri'
    run.font.size = Pt(11)
    run.bold = True
    run.underline = True
    # 5 空行
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    # 6 请从我公司账户中，支付以下2016年8月的工资(包括津贴和相关的调整)，汇款用途为Payroll，合计金额为  582,423.00 美元，并请按下述交易指示办理：
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph_format.left_indent = Pt(4)
    # 6.1 请从我公司账户中，支付以下
    run = paragraph.add_run(u'请从我公司账户中，支付以下')
    run.font.name = u'宋体'
    run.font.size = Pt(10)
    # 6.2 年份
    run = paragraph.add_run(u'%s' % period_year)
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    run.bold = True
    # 6.3 年
    run = paragraph.add_run(u'年')
    run.font.name = u'宋体'
    run.font.size = Pt(10)
    # 6.4 月份
    run = paragraph.add_run(u'%s' % period_month)
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    run.bold = True
    # 6.5 月的工资
    run = paragraph.add_run(u'月的工资')
    run.font.name = u'宋体'
    run.font.size = Pt(10)
    # 6.6 (
    run = paragraph.add_run(u'(')
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    # 6.7 包括津贴和相关的调整
    run = paragraph.add_run(u'包括津贴和相关的调整')
    run.font.name = u'宋体'
    run.font.size = Pt(10)
    # 6.8 )
    run = paragraph.add_run(u')')
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    # 6.9 ，汇款用途为Payroll，合计金额为
    run = paragraph.add_run(u'，汇款用途为Payroll，合计金额为')
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 6.10 金额 '  %s '
    run = paragraph.add_run(u'  %s ' % total_money)
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    run.bold = True
    # 6.11 美元，并请按下述交易指示办理：
    run = paragraph.add_run(u'美元，并请按下述交易指示办理：')
    run.font.name = u'宋体'
    run.font.size = Pt(10)
    # 7 u'Please debit from our account to pay the salaries of %s %s with a total amount of USD %s. Payment Purpose is Payroll; kindly refer to below payment details for further processing:'
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph_format.left_indent = Pt(4)
    # 7.1 Please debit from our account to pay the salaries of
    run = paragraph.add_run(u'Please debit from our account to pay the salaries of')
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    # 7.2 月份 年
    run = paragraph.add_run(u' %s %s' % (MONTH_EN_MAP[str(period_month)], period_year))
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    run.bold = True
    # 7.3  with a total amount of USD
    run = paragraph.add_run(u' with a total amount of USD')
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    # 7.4 金额
    run = paragraph.add_run(u' %s' % total_money)
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    run.bold = True
    # 7.5 . Payment Purpose is Payroll; kindly refer to below payment details for further processing:
    run = paragraph.add_run(u'. Payment Purpose is Payroll; kindly refer to below payment details for further processing:')
    run.font.name = u'Calibri'
    run.font.size = Pt(10)
    pass
    # 字体样式待完善
    # 8 扣帐账户名称: 亚洲基础设施投资银行          		扣帐帐户号码: 012-875-9-277984-8           		扣帐日期: 	30 August 201
    # paragraph = document.add_paragraph(u'扣帐账户名称: 亚洲基础设施投资银行          		扣帐帐户号码: 012-875-9-277984-8           		扣帐日期: 	%s	')
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph_format.left_indent = Pt(8)
    paragraph = document.add_paragraph()
    # 8.1 扣帐账户名称:
    run = paragraph.add_run(u'扣帐账户名称:')
    run.font.name = u'Calibri'
    run.font.size = Pt(11)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 8.2  亚洲基础设施投资银行
    run = paragraph.add_run(u'亚洲基础设施投资银行          ')
    run.font.name = u'宋体'
    run.font.size = Pt(12)
    run.underline = True
    # 8.3 		扣帐帐户号码:
    run = paragraph.add_run(u'		扣帐帐户号码: ')
    run.font.name = u'宋体'
    run.font.size = Pt(11)
    r = run._element
    r.rPr.rFonts.set(qn('w:ascii'), u'Calibri')
    # 8.4 012-875-9-277984-8
    run = paragraph.add_run(u'012-875-9-277984-8           ')
    run.font.name = u'Calibri'
    run.font.size = Pt(12)
    run.underline = True
    # 8.5 		扣帐日期:
    run = paragraph.add_run(u'		扣帐日期: ')
    run.font.name = u'宋体'
    run.font.size = Pt(11)
    r = run._element
    r.rPr.rFonts.set(qn('w:ascii'), u'Calibri')
    # 8.4 	日 月 年
    run = paragraph.add_run(u'	%s %s %s	' % (day, MONTH_EN_MAP[str(month)], year))
    run.font.name = u'Calibri'
    run.font.size = Pt(12)
    run.underline = True
    run.bold = True
    # 9 开始表格制作
    for index, i in enumerate(rule_result_data):
        if index % 10 == 0 and index != 0:
            # 结尾文字 Part0
            paragraph = document.add_paragraph(u'')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)           # 上行间距
            paragraph_format.space_after = Pt(0)            # 下行间距
            paragraph_format.line_spacing = Pt(9)           # 行间距设置
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_format.left_indent = Pt(4)
            # 加椭圆图形
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)           # 上行间距
            paragraph_format.space_after = Pt(0)            # 下行间距
            paragraph_format.line_spacing = Pt(9)           # 行间距设置
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.left_indent = Pt(4)
            run = paragraph.add_run()
            run.add_picture(u'../../ehr-new/ehr/custom_addons/addons/ey_report/static/description/ellipse.png', width=Cm(1.6), height=Cm(1.3))
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # 结尾文字 Part1
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)           # 上行间距
            paragraph_format.space_after = Pt(0)            # 下行间距
            paragraph_format.line_spacing = Pt(9)           # 行间距设置
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.left_indent = Pt(4)
            run = paragraph.add_run(u'_____________________________________________')
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run.font.name = u'Calibri'
            run.font.size = Pt(10.5)
            # 结尾文字 Part2
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)           # 上行间距
            paragraph_format.space_after = Pt(0)            # 下行间距
            paragraph_format.line_spacing = Pt(9)           # 行间距设置
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.left_indent = Pt(4)
            run = paragraph.add_run(u'客户被授权人签署 Authorized Signatory(ies)')
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run.font.name = u'宋体'
            run.font.size = Pt(10.5)
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), u'Calibri')
            # 分页
            document.add_page_break()
            # 页头段落1
            paragraph = document.add_paragraph()
            run = paragraph.add_run(u'传真委办发薪交易指示表格 (续)')
            run.font.name = u'宋体'
            run.font.size = Pt(9)
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), u'Calibri')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.keep_with_next = True
            paragraph_format.space_before = Pt(0)           # 上行间距
            paragraph_format.space_after = Pt(0)            # 下行间距
            paragraph_format.line_spacing = Pt(9)           # 行间距设置
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.left_indent = Pt(4)
            # 页头段落2
            paragraph = document.add_paragraph()
            run = paragraph.add_run(u'Payroll Instruction by Facsimile (Continued)')
            run.font.name = u'Calibri'
            run.font.size = Pt(9)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.keep_with_next = True
            paragraph_format.space_before = Pt(0)           # 上行间距
            paragraph_format.space_after = Pt(0)            # 下行间距
            paragraph_format.line_spacing = Pt(9)           # 行间距设置
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.left_indent = Pt(4)
            # 页头段落3
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)           # 上行间距
            paragraph_format.space_after = Pt(0)            # 下行间距
            paragraph_format.line_spacing = Pt(9)           # 行间距设置
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.left_indent = Pt(4)
            # 页头段落4
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)           # 上行间距
            paragraph_format.space_after = Pt(0)            # 下行间距
            paragraph_format.line_spacing = Pt(9)           # 行间距设置
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph_format.left_indent = Pt(4)
        if index % 10 == 0:
            if index != 0:
                # table.autofit = False
                set_row_width(table.rows)
                set_column_width(table.columns)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table = document.add_table(rows=1, cols=9, style='Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.width = Cm(27.71)
            table.rows[-1].height = Cm(3.23)
            # tr = table.rows[0]._tr
            # trPr = tr.get_or_add_trPr()
            # trHeight = OxmlElement('w:trHeight')
            # w_val = Cm(3.23)
            # trHeight.set(qn('w:val'), str(w_val))
            # trHeight.set(qn('w:hRule'), "atLeast")
            # trPr.append(trHeight)
            # if index == 10:
            #     document.paragraphs[-1].text = ''
            # table.autofit = False
            # [u'name', u'code', u'value_type', u'is_necessary', u'is_only', u'cumulative_item']
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'序号 No.'
            hdr_cells[1].text = u'收款人姓名 Beneficiary Name'
            hdr_cells[2].text = u'支付金额(美元) Amount(USD)'
            hdr_cells[3].text = u'收款账号 Beneficiary Account Number'
            hdr_cells[4].text = u'账户名称 Beneficiary Account Name'
            hdr_cells[5].text = u'收款银行总行／分行／支行名称 Beneficiary Bank’s Branch Name'
            hdr_cells[6].text = u'收款银行名称 Beneficiary Bank Name'
            hdr_cells[7].text = u'收款银行 国家地区 Country Region of the Beneficiary Bank'
            hdr_cells[8].text = u'收款人联络电话 Beneficiary’s Contact Number'
            # table.rows[0].width = Cm(3.23)
            # para = hdr_cells[0].add_paragraph(u'序号 No.')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # para = hdr_cells[1].add_paragraph(u'收款人姓名 Beneficiary Name')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # para = hdr_cells[2].add_paragraph(u'支付金额(美元) Amount(USD)')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # para = hdr_cells[3].add_paragraph(u'收款账号 Beneficiary Account Number')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # para = hdr_cells[4].add_paragraph(u'账户名称 Beneficiary Account Name')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # para = hdr_cells[5].add_paragraph(u'收款银行总行／分行／支行名称 Beneficiary Bank’s Branch Name')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # para = hdr_cells[6].add_paragraph(u'收款银行名称 Beneficiary Bank Name')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # para = hdr_cells[7].add_paragraph(u'收款银行 国家地区 Country Region of the Beneficiary Bank')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # para = hdr_cells[8].add_paragraph(u'收款人联络电话 Beneficiary’s Contact Number')
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_row = table.add_row()
        new_row.height = Cm(0.56)
        # tr = new_row._tr
        # trPr = tr.get_or_add_trPr()
        # trHeight = OxmlElement('w:trHeight')
        # w_val = Cm(0.56)
        # trHeight.set(qn('w:val'), str(w_val))
        # trHeight.set(qn('w:hRule'), 'atLeast')
        # trPr.append(trHeight)
        row_cells = new_row.cells
        table.rows[-1].width = Cm(0.56)
        row_cells[0].text = unicode(index+1)
        for kix, k in enumerate(KEY_LIST):
            row_cells[kix+1].text = unicode(i[k])
            # para = row_cells[kix+1].add_paragraph(unicode(i[k]))
            # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ending Part0
    # 结尾文字 Part1
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph_format.left_indent = Pt(4)
    run = paragraph.add_run(u'_____________________________________________')
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run.font.name = u'Calibri'
    run.font.size = Pt(10.5)
    print paragraph_format.alignment
    # 结尾文字 Part2
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph_format.left_indent = Pt(4)
    run = paragraph.add_run(u'客户被授权人签署 Authorized Signatory(ies)')
    run.font.name = u'宋体'
    run.font.size = Pt(10.5)
    r = run._element
    r.rPr.rFonts.set(qn('w:ascii'), u'Calibri')
    print paragraph_format.alignment
    # Ending Part1
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph_format.left_indent = Pt(4)
    run = paragraph.add_run(u'***填妥此委办交易指示表格后，请实时传真至本行 (传真号码 852-3406-2301)***')
    run.font.name = u'宋体'
    run.font.size = Pt(10.5)
    r = run._element
    r.rPr.rFonts.set(qn('w:ascii'), u'Calibri')

    # Ending Part2
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)           # 上行间距
    paragraph_format.space_after = Pt(0)            # 下行间距
    paragraph_format.line_spacing = Pt(9)           # 行间距设置
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph_format.left_indent = Pt(4)
    run = paragraph.add_run(u'***Upon the completion of the instruction, please fax to us at 852 3406 2301 for processing ***')
    run.font.name = u'宋体'
    run.font.size = Pt(9)
    r = run._element
    r.rPr.rFonts.set(qn('w:ascii'), u'Calibri')

    # 整理页脚问题
    # styles = document.styles
    # footer = styles[WD_STYLE.FOOTER]
    # sections = document.sections
    # total_pages = len(sections)
    # for idx, m_s in enumerate(sections):
    #     m_s.header.text = '%d of %d | Page' % (idx+1, total_pages)
    #     m_s.header.is_linked_to_previous = True

    #保存文件
    document.save(u'测试4.docx')
    # fp = StringIO()
    # document.save(fp)
    # fp.seek(0)
    # docx_file = fp.read()
    # fp.close()
    # return docx_file

def create_docx_file():
    data = get_excel_data()[u'工作表1']
    make_docx_file_format('2016', '11', 123456, data)

def open_docx_file():
    f = open('/Users/dongsk/Downloads/HKBankFileSample.docx')
    print f


import zipfile
import shutil as su
import os
import tempfile
import xml.etree.cElementTree


def get_word_xml(docx_filename):
   with open(docx_filename, mode='rt') as f:
      zip = zipfile.ZipFile(f)
      xml_content = zip.read('word/document.xml')
   return xml_content

def write_and_close_docx (self, xml_content, output_filename):
        """ Create a temp directory, expand the original docx zip.
            Write the modified xml to word/document.xml
            Zip it up as the new docx
        """

        tmp_dir = tempfile.mkdtemp()

        self.zipfile.extractall(tmp_dir)

        with open(os.path.join(tmp_dir,'word/document.xml'), 'w') as f:
            xmlstr = tree.tostring(xml_content, pretty_print=True)
            f.write(xmlstr)

        # Get a list of all the files in the original docx zipfile
        filenames = self.zipfile.namelist()
        # Now, create the new zip file and add all the filex into the archive
        zip_copy_filename = output_filename
        with zipfile.ZipFile(zip_copy_filename, "w") as docx:
            for filename in filenames:
                docx.write(os.path.join(tmp_dir,filename), filename)

        # Clean up the temp dir
        su.rmtree(tmp_dir)

def get_xml_tree(f):
    return xml.etree.ElementTree.parse(f)


if __name__ == '__main__':
    create_docx_file()
    # open_docx_file()
    # word_doc = '/Users/dongsk/Downloads/HKBankFileSample.docx'
    # new_word_doc = 'SLIM.docx'
    # doc = get_word_xml(word_doc)
    # tree = get_xml_tree(doc)
    # write_and_close_docx(word_doc, tree, new_word_doc)