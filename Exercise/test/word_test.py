#coding=utf-8

import xlrd
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Inches, Cm
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING

MY_ELLIPSIS = u"""<mc:AlternateContent>
    <mc:Choice Requires="wps">
      <w:drawing>
        <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251659264" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" wp14:anchorId="3840C1B4" wp14:editId="2753EB23">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="column">
            <wp:posOffset>9344025</wp:posOffset>
          </wp:positionH>
          <wp:positionV relativeFrom="paragraph">
            <wp:posOffset>36195</wp:posOffset>
          </wp:positionV>
          <wp:extent cx="571500" cy="457200"/>
          <wp:effectExtent l="0" t="0" r="19050" b="19050"/>
          <wp:wrapNone/>
          <wp:docPr id="1" name="Oval 2"/>
          <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"/>
          </wp:cNvGraphicFramePr>
          <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:wsp>
                <wps:cNvSpPr>
                  <a:spLocks noChangeArrowheads="1"/>
                </wps:cNvSpPr>
                <wps:spPr bwMode="auto">
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="571500" cy="457200"/>
                  </a:xfrm>
                  <a:prstGeom prst="ellipse">
                    <a:avLst/>
                  </a:prstGeom>
                  <a:solidFill>
                    <a:srgbClr val="FFFFFF"/>
                  </a:solidFill>
                  <a:ln w="9525">
                    <a:solidFill>
                      <a:srgbClr val="000000"/>
                    </a:solidFill>
                    <a:round/>
                    <a:headEnd/>
                    <a:tailEnd/>
                  </a:ln>
                </wps:spPr>
                <wps:txbx>
                  <w:txbxContent>
                    <w:p w:rsidR="001211A4" w:rsidRDefault="001211A4" w:rsidP="001211A4">
                      <w:pPr>
                        <w:snapToGrid w:val="0"/>
                        <w:jc w:val="center"/>
                        <w:rPr>
                          <w:b/>
                          <w:bCs/>
                          <w:sz w:val="22"/>
                        </w:rPr>
                      </w:pPr>
                      <w:r w:rsidRPr="00B25095">
                        <w:rPr>
                          <w:rFonts w:eastAsia="宋体"/>
                          <w:b/>
                          <w:bCs/>
                          <w:sz w:val="22"/>
                          <w:lang w:eastAsia="zh-CN"/>
                        </w:rPr>
                        <w:t>s.v.</w:t>
                      </w:r>
                    </w:p>
                    <w:p w:rsidR="001211A4" w:rsidRDefault="001211A4" w:rsidP="001211A4">
                      <w:pPr>
                        <w:snapToGrid w:val="0"/>
                      </w:pPr>
                    </w:p>
                  </w:txbxContent>
                </wps:txbx>
                <wps:bodyPr rot="0" vert="horz" wrap="square" lIns="91440" tIns="45720" rIns="91440" bIns="45720" anchor="t" anchorCtr="0" upright="1">
                  <a:noAutofit/>
                </wps:bodyPr>
              </wps:wsp>
            </a:graphicData>
          </a:graphic>
          <wp14:sizeRelH relativeFrom="page">
            <wp14:pctWidth>0</wp14:pctWidth>
          </wp14:sizeRelH>
          <wp14:sizeRelV relativeFrom="page">
            <wp14:pctHeight>0</wp14:pctHeight>
          </wp14:sizeRelV>
        </wp:anchor>
      </w:drawing>
    </mc:Choice>
    <mc:Fallback>
      <w:pict>
        <v:oval id="Oval 2" o:spid="_x0000_s1026" style="position:absolute;margin-left:735.75pt;margin-top:2.85pt;width:45pt;height:36pt;z-index:251659264;visibility:visible;mso-wrap-style:square;mso-width-percent:0;mso-height-percent:0;mso-wrap-distance-left:9pt;mso-wrap-distance-top:0;mso-wrap-distance-right:9pt;mso-wrap-distance-bottom:0;mso-position-horizontal:absolute;mso-position-horizontal-relative:text;mso-position-vertical:absolute;mso-position-vertical-relative:text;mso-width-percent:0;mso-height-percent:0;mso-width-relative:page;mso-height-relative:page;v-text-anchor:top" o:gfxdata="UEsDBBQABgAIAAAAIQDkmcPA+wAAAOEBAAATAAAAW0NvbnRlbnRfVHlwZXNdLnhtbJSRQU7DMBBF&#10;90jcwfIWJQ5dIISSdEHaJSBUDjCyJ4nVZGx53NDeHictG4SKWNrj9//TuFwfx0FMGNg6quR9XkiB&#10;pJ2x1FXyY7fNHqXgCGRgcISVPCHLdX17U+5OHlkkmriSfYz+SSnWPY7AufNIadK6MEJMx9ApD3oP&#10;HapVUTwo7SgixSzOGbIuG2zhMESxOabrs0nCpXg+v5urKgneD1ZDTKJqnqpfuYADXwEnMj/ssotZ&#10;nsglnHvr+e7S8JpWE6xB8QYhvsCYPJQJrHDlGqfz65Zz2ciZa1urMW8Cbxbqr2zjPing9N/wJmHv&#10;OH2nq+WD6i8AAAD//wMAUEsDBBQABgAIAAAAIQAjsmrh1wAAAJQBAAALAAAAX3JlbHMvLnJlbHOk&#10;kMFqwzAMhu+DvYPRfXGawxijTi+j0GvpHsDYimMaW0Yy2fr28w6DZfS2o36h7xP//vCZFrUiS6Rs&#10;YNf1oDA78jEHA++X49MLKKk2e7tQRgM3FDiMjw/7My62tiOZYxHVKFkMzLWWV63FzZisdFQwt81E&#10;nGxtIwddrLvagHro+2fNvxkwbpjq5A3wye9AXW6lmf+wU3RMQlPtHCVN0xTdPaoObMsc3ZFtwjdy&#10;jWY5YDXgWTQO1LKu/Qj6vn74p97TRz7jutV+h4zrj1dvuhy/AAAA//8DAFBLAwQUAAYACAAAACEA&#10;OH4CvRwCAAA+BAAADgAAAGRycy9lMm9Eb2MueG1srFNRb9MwEH5H4j9YfqdJq5axqOk0dRQhDTZp&#10;8AMcx2ksHJ85u03Kr+fsZKUDnhB5sO5y589333e3vhk6w44KvQZb8vks50xZCbW2+5J//bJ7844z&#10;H4SthQGrSn5Snt9sXr9a965QC2jB1AoZgVhf9K7kbQiuyDIvW9UJPwOnLAUbwE4EcnGf1Sh6Qu9M&#10;tsjzt1kPWDsEqbynv3djkG8SftMoGR6axqvATMmptpBOTGcVz2yzFsUehWu1nMoQ/1BFJ7SlR89Q&#10;dyIIdkD9B1SnJYKHJswkdBk0jZYq9UDdzPPfunlqhVOpFyLHuzNN/v/Bys/HR2S6Ju04s6IjiR6O&#10;wrBFZKZ3vqCEJ/eIsTfv7kF+88zCthV2r24RoW+VqKmeeczPXlyIjqerrOo/QU3A4hAgkTQ02EVA&#10;ap8NSYvTWQs1BCbp5+pqvspJMUmh5eqKtE4viOL5skMfPijoWDRKrozRzke2RCGO9z7EekTxnJXq&#10;B6PrnTYmObivtgYZNVvyXfqmB/xlmrGsL/n1arFKyC9i/hIiT9/fIBAOtk5zFrl6P9lBaDPaVKWx&#10;E3mRr5H3MFTDpMukRAX1idhEGIeYlo6MFvAHZz0NcMn994NAxZn5aEmR6/lyGSc+OYlBzvAyUl1G&#10;hJUEVfLA2Whuw7glB4d639JL80SAhVtSsdGJ3qjwWNVUPg1pYn1aqLgFl37K+rX2m58AAAD//wMA&#10;UEsDBBQABgAIAAAAIQC3U32j3QAAAAoBAAAPAAAAZHJzL2Rvd25yZXYueG1sTI/BToQwEIbvJr5D&#10;Mybe3IJYapCy2bgx0YMHUe9dOgtk6ZTQLotvbznp8Z/58s835XaxA5tx8r0jBekmAYbUONNTq+Dr&#10;8+XuEZgPmoweHKGCH/Swra6vSl0Yd6EPnOvQslhCvtAKuhDGgnPfdGi137gRKe6ObrI6xDi13Ez6&#10;EsvtwO+TJOdW9xQvdHrE5w6bU322Cvbtrs5nngWRHfevQZy+39+yVKnbm2X3BCzgEv5gWPWjOlTR&#10;6eDOZDwbYn6QqYisAiGBrYDI18FBgZQSeFXy/y9UvwAAAP//AwBQSwECLQAUAAYACAAAACEA5JnD&#10;wPsAAADhAQAAEwAAAAAAAAAAAAAAAAAAAAAAW0NvbnRlbnRfVHlwZXNdLnhtbFBLAQItABQABgAI&#10;AAAAIQAjsmrh1wAAAJQBAAALAAAAAAAAAAAAAAAAACwBAABfcmVscy8ucmVsc1BLAQItABQABgAI&#10;AAAAIQA4fgK9HAIAAD4EAAAOAAAAAAAAAAAAAAAAACwCAABkcnMvZTJvRG9jLnhtbFBLAQItABQA&#10;BgAIAAAAIQC3U32j3QAAAAoBAAAPAAAAAAAAAAAAAAAAAHQEAABkcnMvZG93bnJldi54bWxQSwUG&#10;AAAAAAQABADzAAAAfgUAAAAA&#10;">
          <v:textbox>
            <w:txbxContent>
              <w:p w:rsidR="001211A4" w:rsidRDefault="001211A4" w:rsidP="001211A4">
                <w:pPr>
                  <w:snapToGrid w:val="0"/>
                  <w:jc w:val="center"/>
                  <w:rPr>
                    <w:b/>
                    <w:bCs/>
                    <w:sz w:val="22"/>
                  </w:rPr>
                </w:pPr>
                <w:r w:rsidRPr="00B25095">
                  <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:bCs/>
                    <w:sz w:val="22"/>
                    <w:lang w:eastAsia="zh-CN"/>
                  </w:rPr>
                  <w:t>s.v.</w:t>
                </w:r>
              </w:p>
              <w:p w:rsidR="001211A4" w:rsidRDefault="001211A4" w:rsidP="001211A4">
                <w:pPr>
                  <w:snapToGrid w:val="0"/>
                </w:pPr>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:oval>
      </w:pict>
    </mc:Fallback>
  </mc:AlternateContent>"""

def get_excel_data():
    data = {}
    file_name = '/Users/dongsk/Desktop/工作簿1.xlsx'
    # file_name = '/Users/dongsk/Downloads/1静态数据准备.xlsx'
    file_data = xlrd.open_workbook(file_name)
    tables = file_data.sheets()
    tables = file_data.sheets()
    for table in tables:
        print "开始处理 ",table.name
        nrows = table.nrows    # 行数
        if nrows <= 1:
            print "文件 < %s > 没有内容" % (table.name)
            continue
        ncols = table.ncols     # 列数
        colnames = table.row_values(0)  # 第一行数据
        leave_none_list = []
        data[table.name] = []
        for rownum in range(1, nrows):
            data_dict = {}
            row = table.row_values(rownum)
            if row:
                for i in range(len(colnames)):
                    ccol = colnames[i]
                    rrow = row[i]
                    data_dict[ccol] = rrow
                data[table.name].append(data_dict)
        print table.name, "处理完成"

    return data

#打开文档
document = Document()
sections = document.sections
print len(sections)
# 设置格式
# section = document.add_section()
section = document.sections[0]
# section.page_width = Inches(11.69)
# section.page_height = Inches(8.27)
section.orientation = WD_ORIENT.LANDSCAPE
section.page_height = Cm(20.99)
section.page_width = Cm(29.7)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.393)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
# section = document.sections[-1]
# PORTRAIT 纵向
# section.orientation = WD_ORIENT.PORTRAIT
# LANDSCAPE 横向
# section.orientation = WD_ORIENT.LANDSCAPE

#加入不同等级的标题
document.add_heading(u'MS WORD写入测试',0)
document.add_heading(u'一级标题',1)
document.add_heading(u'二级标题',2)
#添加文本
paragraph = document.add_paragraph(u'我们在做文本测试！')
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(0)           # 上行间距
paragraph_format.space_after = Pt(0)            # 下行间距
# paragraph_format.line_spacing = Pt(9)           # 行间距设置
paragraph.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # EXACTLY
#设置字号
run = paragraph.add_run(u'设置字号、')
pic = run.add_picture(u'/Users/dongsk/Downloads/图片/董.jpg', width=Cm(1.6), height=Cm(1.3))

run.font.size = Pt(24)

#设置字体
run = paragraph.add_run('Set Font,')
run.font.name = 'Consolas'

#设置中文字体
run = paragraph.add_run(u'设置中文字体、payroll哈哈,stop下来之后restare')
run.font.name=u'Calibri'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

#设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True

#设置粗体
run = paragraph.add_run(u'粗体').bold = True

#增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

#增加无序列表
document.add_paragraph(
    u'无序列表元素1', style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2', style='List Bullet'
)
#增加有序列表
document.add_paragraph(
    u'有序列表元素1', style='List Number'
)
document.add_paragraph(
    u'有序列表元素2', style='List Number'
)
#增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
document.add_picture(u'/Users/dongsk/Desktop/5倍大图椭圆.png', width=Cm(1.6), height=Cm(1.3))

#增加表格
table = document.add_table(rows=1, cols=3)
# table.style.hidden = True
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
#再增加3行表格元素
for i in xrange(3):
    row_cells = table.add_row().cells
    row_cells[0].text = 'test'+str(i)
    row_cells[1].text = str(i)
    row_cells[2].text = 'desc'+str(i)

#增加分页
document.add_page_break()
paragraph1 = document.add_paragraph()
run1 = paragraph1.add_run(u'传真委办发薪交易指示表格 (续)')
paragraph1.paragraph_format.keep_with_next = True

document.add_page_break()
paragraph = document.add_paragraph()
doc2 = docx.Document('一个椭圆.docx')
run_esp = doc2.paragraphs[0].runs[0]
run = paragraph.add_run()
run_esp._parent = run._parent
run._element = run_esp._element
# paragraph.runs.append(run_esp)
# run = paragraph.add_run()
# pic = run.add_picture()
# print '\n&&&&&&&&&&&\n', run._element.xml, '\n&&&&&&&&&&&\n', type(run._element.xml), '\n&&&&&&&&&&\n'
# abc = run._element.xml.replace('</w:rPr>', '</w:rPr>' + MY_ELLIPSIS)
# print '\n&&&&&&&&&&&\n', abc, '\n&&&&&&&&&&&\n', type(abc), '\n&&&&&&&&&&\n'
# run._element.xml += abc
# run._element.xml += MY_ELLIPSIS



document.add_page_break()



# section = document.sections[0]
# PORTRAIT 纵向
# section.orientation = WD_ORIENT.PORTRAIT
# LANDSCAPE 横向
# section.orientation = WD_ORIENT.LANDSCAPE

#保存文件
document.save(u'测试1.docx')

#
#
# from docx import Document


# document = Document()
# # paragraph_format.line_spacing = Pt(9)
# paragraph = document.add_paragraph()
# # paragraph.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
# paragraph_format = paragraph.paragraph_format
# paragraph_format.space_before = Pt(0)           # 上行间距
# paragraph_format.space_after = Pt(0)            # 下行间距
# paragraph_format.line_spacing = Pt(9)           # 行间距设置
# paragraph.line_spacing_rule = WD_LINE_SPACING.AT_LEAST(Pt(9))
# # paragraph_format.line_spacing_rule = 'min'      # 行间距设置

styles = document.styles
table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]
for style in table_styles:
    print(style.name)

